# Simple Markdown to Word Converter - QGIS Compatible
# Version 1.2 - Minimal dependencies
# Run this in QGIS Python Console

import os

def convert_md_to_word():
    """Basic markdown to word converter with minimal dependencies"""
    
    print("=== Markdown to Word Converter v1.2 ===")
    
    # Check if we can import required modules
    try:
        from qgis.PyQt.QtWidgets import QFileDialog, QMessageBox
        print("✓ QGIS GUI modules loaded")
    except ImportError as e:
        print(f"✗ Cannot load QGIS modules: {e}")
        return False
    
    try:
        from docx import Document
        print("✓ python-docx module loaded")
    except ImportError:
        QMessageBox.critical(None, "Missing Package", 
                           "python-docx not found!\n\nInstall with:\npip install python-docx")
        return False
    
    # File selection
    md_files, _ = QFileDialog.getOpenFileNames(
        None,
        "Select Markdown Files",
        "",
        "Markdown Files (*.md *.markdown);;All Files (*)"
    )
    
    if not md_files:
        print("No files selected")
        return False
    
    # Output directory selection
    output_dir = QFileDialog.getExistingDirectory(
        None,
        "Select Output Directory",
        os.path.expanduser("~/Documents")
    )
    
    if not output_dir:
        print("No output directory selected")
        return False
    
    print(f"Processing {len(md_files)} files to: {output_dir}")
    
    # Process files
    success_count = 0
    
    for md_file in md_files:
        try:
            print(f"Processing: {os.path.basename(md_file)}")
            
            # Read markdown file
            with open(md_file, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            # Create Word document
            doc = Document()
            
            # Basic processing - split by lines
            lines = content.split('\n')
            
            for line in lines:
                line = line.strip()
                
                if not line:
                    continue
                elif line.startswith('# '):
                    # Heading 1
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    # Heading 2
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    # Heading 3
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('#### '):
                    # Heading 4
                    doc.add_heading(line[5:], level=4)
                elif line.startswith('- ') or line.startswith('* '):
                    # Bullet point
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
                    # Numbered list
                    doc.add_paragraph(line[3:], style='List Number')
                else:
                    # Regular paragraph
                    doc.add_paragraph(line)
            
            # Save document
            base_name = os.path.splitext(os.path.basename(md_file))[0]
            output_file = os.path.join(output_dir, f"{base_name}.docx")
            doc.save(output_file)
            
            print(f"  ✓ Saved: {base_name}.docx")
            success_count += 1
            
        except Exception as e:
            print(f"  ✗ Error processing {os.path.basename(md_file)}: {e}")
            continue
    
    # Show results
    message = f"Conversion completed!\n\nProcessed: {success_count}/{len(md_files)} files\nOutput: {output_dir}"
    print("\n" + "="*50)
    print(message)
    
    try:
        QMessageBox.information(None, "Conversion Complete", message)
    except:
        pass
    
    return True

# Try to run immediately if in QGIS
try:
    from qgis.core import QgsApplication
    if QgsApplication.instance():
        print("Ready to convert! Run: convert_md_to_word()")
        print("Or auto-starting conversion...")
        convert_md_to_word()
except Exception as e:
    print(f"QGIS check failed: {e}")
    print("Run convert_md_to_word() manually")
